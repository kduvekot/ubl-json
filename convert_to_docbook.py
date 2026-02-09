#!/usr/bin/env python3
"""Convert 'UBL JSON Syntax Binding.docx' to DocBook XML 4.5."""

import re
import docx
from lxml import etree
import zipfile

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W = f"{{{WNS}}}"
R = f"{{{RNS}}}"

DOCX_PATH = "UBL JSON Syntax Binding.docx"
OUTPUT_PATH = "UBL-2.5-JSON-Syntax-Binding.xml"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def slugify(text):
    """Create an uppercase slug from text for use in XML ids."""
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text.strip())
    return text.upper()


def xml_escape(text):
    """Escape special XML characters."""
    if text is None:
        return ""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


# ---------------------------------------------------------------------------
# Build hyperlink relationship map
# ---------------------------------------------------------------------------

def build_hyperlink_map(doc):
    """Map relationship IDs to URLs for external hyperlinks."""
    rels = doc.part.rels
    hmap = {}
    for rel_id, rel in rels.items():
        if "hyperlink" in str(rel.reltype).lower():
            hmap[rel_id] = rel.target_ref
    return hmap


# ---------------------------------------------------------------------------
# Build numbering format map
# ---------------------------------------------------------------------------

def build_numbering_map(docx_path):
    """Build a map: (numId, ilvl) -> 'bullet' | 'decimal' | ... """
    with zipfile.ZipFile(docx_path) as z:
        with z.open("word/numbering.xml") as f:
            tree = etree.fromstring(f.read())

    ns = {"w": WNS}
    # abstractNumId -> { ilvl -> numFmt }
    abstract_map = {}
    for an in tree.findall(".//w:abstractNum", ns):
        aid = an.get(f"{W}abstractNumId")
        abstract_map[aid] = {}
        for lvl in an.findall(".//w:lvl", ns):
            ilvl = lvl.get(f"{W}ilvl")
            nf = lvl.find(".//w:numFmt", ns)
            fmt = nf.get(f"{W}val") if nf is not None else "bullet"
            abstract_map[aid][ilvl] = fmt

    # numId -> abstractNumId
    num_to_abstract = {}
    for num in tree.findall(".//w:num", ns):
        nid = num.get(f"{W}numId")
        anr = num.find(".//w:abstractNumId", ns)
        if anr is not None:
            num_to_abstract[nid] = anr.get(f"{W}val")

    # (numId, ilvl) -> fmt
    result = {}
    for nid, aid in num_to_abstract.items():
        if aid in abstract_map:
            for ilvl, fmt in abstract_map[aid].items():
                result[(nid, ilvl)] = fmt
    return result


# ---------------------------------------------------------------------------
# Inline content extraction
# ---------------------------------------------------------------------------

def extract_inline_content(paragraph, hyperlink_map):
    """Walk paragraph XML to extract runs and hyperlinks in document order.

    Returns a list of dicts:
      {'type': 'text', 'text': str, 'bold': bool, 'italic': bool}
      {'type': 'hyperlink', 'url': str, 'text': str}
    """
    result = []
    p_elem = paragraph._element

    for child in p_elem:
        tag = child.tag
        if tag == f"{W}r":
            # Collect text from <w:t> and treat <w:tab/> as a tab character
            text_parts = []
            for sub in child:
                if sub.tag == f"{W}t" and sub.text:
                    text_parts.append(sub.text)
                elif sub.tag == f"{W}tab":
                    text_parts.append("\t")
            text = "".join(text_parts)
            if not text:
                continue
            rPr = child.find(f"{W}rPr")
            bold = False
            italic = False
            if rPr is not None:
                b_elem = rPr.find(f"{W}b")
                if b_elem is not None:
                    val = b_elem.get(f"{W}val")
                    bold = val != "0" if val is not None else True
                i_elem = rPr.find(f"{W}i")
                if i_elem is not None:
                    val = i_elem.get(f"{W}val")
                    italic = val != "0" if val is not None else True
            result.append({"type": "text", "text": text, "bold": bold, "italic": italic})

        elif tag == f"{W}hyperlink":
            rid = child.get(f"{R}id")
            texts = []
            for r in child.findall(f".//{W}t"):
                if r.text:
                    texts.append(r.text)
            link_text = "".join(texts)
            url = hyperlink_map.get(rid, "") if rid else ""
            if link_text:
                result.append({"type": "hyperlink", "url": url, "text": link_text})

    return result


def render_inline(items):
    """Render a list of inline items to DocBook XML string."""
    parts = []
    for item in items:
        if item["type"] == "hyperlink":
            url = xml_escape(item["url"])
            text = xml_escape(item["text"])
            if url:
                parts.append(f'<ulink url="{url}">{text}</ulink>')
            else:
                parts.append(text)
        else:
            # Replace tabs with spaces for inline text (tabs are structural
            # separators in the .docx, not meaningful in DocBook inline content)
            text = xml_escape(item["text"].replace("\t", " "))
            if item["bold"] and item["italic"]:
                parts.append(f'<emphasis role="bold"><emphasis>{text}</emphasis></emphasis>')
            elif item["bold"]:
                parts.append(f'<emphasis role="bold">{text}</emphasis>')
            elif item["italic"]:
                parts.append(f'<emphasis>{text}</emphasis>')
            else:
                parts.append(text)
    return "".join(parts)


# ---------------------------------------------------------------------------
# Paragraph info extraction
# ---------------------------------------------------------------------------

def get_paragraph_info(p):
    """Extract style, list info from a paragraph."""
    style = p.style.name if p.style else "Normal"
    numId = None
    ilvl = None
    pPr = p._element.find(f"{W}pPr")
    if pPr is not None:
        numPr = pPr.find(f"{W}numPr")
        if numPr is not None:
            il = numPr.find(f"{W}ilvl")
            ni = numPr.find(f"{W}numId")
            if il is not None:
                ilvl = il.get(f"{W}val")
            if ni is not None:
                numId = ni.get(f"{W}val")
    return style, numId, ilvl


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def render_table(table, hyperlink_map):
    """Render a docx table to DocBook CALS table XML."""
    rows = table.rows
    if not rows:
        return ""

    ncols = len(table.columns)
    lines = []
    lines.append(f'<table id="T-UBL-UNQUALIFIED-DATA-TYPES" frame="all">')
    lines.append(f'  <title>UBL unqualified data types</title>')
    lines.append(f'  <tgroup cols="{ncols}">')
    for ci in range(ncols):
        lines.append(f'    <colspec colnum="{ci+1}" colname="col{ci+1}"/>')

    # First row is header
    lines.append("    <thead>")
    lines.append("      <row>")
    for cell in rows[0].cells:
        lines.append(f"        <entry>{xml_escape(cell.text.strip())}</entry>")
    lines.append("      </row>")
    lines.append("    </thead>")

    lines.append("    <tbody>")
    for row in rows[1:]:
        lines.append("      <row>")
        for cell in row.cells:
            lines.append(f"        <entry>{xml_escape(cell.text.strip())}</entry>")
        lines.append("      </row>")
    lines.append("    </tbody>")
    lines.append("  </tgroup>")
    lines.append("</table>")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert():
    doc = docx.Document(DOCX_PATH)
    hyperlink_map = build_hyperlink_map(doc)
    numbering_map = build_numbering_map(DOCX_PATH)

    # -----------------------------------------------------------------------
    # Phase 1: Parse all paragraphs into structured elements
    # -----------------------------------------------------------------------
    elements = []  # list of dicts describing each document element

    # Track where the table appears -- it's between paragraphs. We need to
    # insert tables by scanning the document body XML for table elements.
    # For simplicity, since there's only one table and it appears in section
    # 10.1 (between the paragraph about data types and the BinaryObject heading),
    # we'll insert it when we detect the right context.

    # Find the table position by scanning for when style goes from Normal with
    # data-types text to "BinaryObject type" heading
    table_inserted = False
    table_insert_after_para = 179  # paragraph index after which the table goes

    for i, p in enumerate(doc.paragraphs):
        style, numId, ilvl = get_paragraph_info(p)

        # Skip ToC entries
        if style.startswith("toc"):
            continue

        # Skip the "Table of Contents" heading
        if style == "Heading 1" and p.text.strip() == "Table of Contents":
            continue

        text = p.text.strip()
        raw_text = p.text  # preserve tabs for definition detection

        inline = extract_inline_content(p, hyperlink_map)

        # Skip truly empty Normal paragraphs
        if not text and not inline and style == "Normal":
            continue

        elem = {
            "index": i,
            "style": style,
            "text": text,
            "raw_text": raw_text,
            "inline": inline,
            "numId": numId,
            "ilvl": ilvl,
        }
        elements.append(elem)

        # Insert table marker after paragraph 179
        if i == table_insert_after_para and not table_inserted:
            elements.append({"style": "TABLE", "index": -1})
            table_inserted = True

    # -----------------------------------------------------------------------
    # Phase 2: Group elements into a tree structure
    # -----------------------------------------------------------------------
    # We'll generate the XML by walking through elements and tracking
    # the section nesting stack.

    xml_lines = []

    def indent(level):
        return "  " * level

    # XML declaration and DOCTYPE
    xml_lines.append('<?xml version="1.0" encoding="UTF-8"?>')
    xml_lines.append('<!DOCTYPE article PUBLIC "-//OASIS//DTD DocBook XML V4.5//EN"')
    xml_lines.append('  "http://www.oasis-open.org/docbook/xml/4.5/docbookx.dtd">')
    xml_lines.append("")

    xml_lines.append('<article lang="en">')
    xml_lines.append("  <articleinfo>")
    xml_lines.append("    <title>UBL 2.5 JSON Syntax Binding</title>")
    xml_lines.append("    <authorgroup>")
    xml_lines.append("      <editor>")
    xml_lines.append("        <firstname>Kenneth</firstname>")
    xml_lines.append("        <surname>Bengtsson</surname>")
    xml_lines.append("      </editor>")
    xml_lines.append("    </authorgroup>")
    xml_lines.append("    <pubdate>2025-09-23</pubdate>")
    xml_lines.append("    <releaseinfo>Revision 306</releaseinfo>")
    xml_lines.append("    <abstract>")
    xml_lines.append("      <para>This document specifies a normative JSON syntax binding for the")
    xml_lines.append("        OASIS Universal Business Language (UBL). It defines the representational")
    xml_lines.append("        rules, structural constraints, and validation requirements by which UBL")
    xml_lines.append("        business documents and their constituent components are to be expressed")
    xml_lines.append("        and verified in JSON.</para>")
    xml_lines.append("    </abstract>")
    xml_lines.append("  </articleinfo>")
    xml_lines.append("")

    # -----------------------------------------------------------------------
    # State machine for generating nested sections
    # -----------------------------------------------------------------------
    heading_level_map = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
    }

    # Section stack: list of heading levels currently open
    section_stack = []
    # Track parent section titles for disambiguating IDs
    section_title_stack = []
    in_bibliography = False
    bib_subsection = None  # "normative" or "informative"
    in_appendix = False
    appendix_name = None
    used_ids = set()

    def close_sections_to(target_level, lines, base_indent=1):
        """Close sections until we're at a level where we can open target_level."""
        while section_stack and section_stack[-1] >= target_level:
            section_stack.pop()
            lvl = base_indent + len(section_stack)
            lines.append(f"{indent(lvl)}</section>")

    def close_all_sections(lines, base_indent=1):
        while section_stack:
            section_stack.pop()
            lvl = base_indent + len(section_stack)
            lines.append(f"{indent(lvl)}</section>")

    # Process elements
    i = 0
    while i < len(elements):
        elem = elements[i]
        style = elem["style"]

        # Handle table insertion marker
        if style == "TABLE":
            tbl = doc.tables[0]
            base = 1 + len(section_stack)
            table_xml = render_table(tbl, hyperlink_map)
            for line in table_xml.split("\n"):
                xml_lines.append(f"{indent(base)}{line}")
            i += 1
            continue

        text = elem.get("text", "")

        # Detect Annex A and Annex B
        if style == "Heading 1" and text.startswith("Annex A"):
            # Close all current sections
            close_all_sections(xml_lines)

            # Annex A as appendix
            xml_lines.append("")
            xml_lines.append(f'  <appendix id="A-LICENSE-DOCUMENT-STATUS-AND-NOTICES">')
            xml_lines.append(f"    <title>{xml_escape(text)}</title>")
            in_appendix = True
            appendix_name = "A"
            section_stack.clear()
            i += 1
            continue

        if style == "Heading 1" and text.startswith("Annex B"):
            # Close appendix A if open
            if in_appendix:
                close_all_sections(xml_lines, base_indent=2)
                xml_lines.append("  </appendix>")
                in_appendix = False

            # Annex B as bibliography appendix
            xml_lines.append("")
            xml_lines.append(f'  <appendix id="A-REFERENCES">')
            xml_lines.append(f"    <title>{xml_escape(text)}</title>")
            in_appendix = True
            in_bibliography = True
            appendix_name = "B"
            section_stack.clear()
            i += 1
            continue

        # Handle headings
        if style in heading_level_map:
            hlevel = heading_level_map[style]

            if in_appendix:
                base = 2
            else:
                base = 1

            # Close sections as needed
            close_sections_to(hlevel, xml_lines, base_indent=base)

            # Also trim the section_title_stack to match
            while len(section_title_stack) >= hlevel:
                section_title_stack.pop()

            # Generate a unique section ID
            slug = slugify(text)
            if not slug:
                slug = "PLACEHOLDER"
            section_id = "S-" + slug
            # If the ID is already used, prefix with parent section title
            if section_id in used_ids:
                parent = section_title_stack[-1] if section_title_stack else ""
                section_id = "S-" + slugify(parent) + "-" + slug
            # Final dedup with numeric suffix if still colliding
            base_id = section_id
            counter = 2
            while section_id in used_ids:
                section_id = f"{base_id}-{counter}"
                counter += 1
            used_ids.add(section_id)

            section_title_stack.append(text)

            # Special handling for Annex B subsections
            if in_bibliography and hlevel == 2:
                if "Normative" in text:
                    bib_subsection = "normative"
                elif "Informative" in text:
                    bib_subsection = "informative"

            lvl = base + len(section_stack)
            xml_lines.append(f"")
            xml_lines.append(f'{indent(lvl)}<section id="{xml_escape(section_id)}">')
            xml_lines.append(f"{indent(lvl+1)}<title>{xml_escape(text)}</title>")
            section_stack.append(hlevel)
            i += 1
            continue

        # Handle Monospace (code blocks) - collect consecutive ones
        if style == "Monospace":
            code_lines = []
            while i < len(elements) and elements[i]["style"] == "Monospace":
                code_lines.append(elements[i]["text"])
                i += 1
            code_text = "\n".join(code_lines)
            lvl = (2 if in_appendix else 1) + len(section_stack)
            xml_lines.append(f"{indent(lvl)}<programlisting><![CDATA[{code_text}]]></programlisting>")
            continue

        # Handle List Paragraphs - collect consecutive ones into a list
        if style == "List Paragraph":
            list_items = []
            while i < len(elements) and elements[i]["style"] == "List Paragraph":
                e = elements[i]
                list_items.append({
                    "text": e["text"],
                    "inline": e["inline"],
                    "numId": e["numId"],
                    "ilvl": int(e["ilvl"] or "0"),
                })
                i += 1

            # Determine if this is an ordered or itemized list
            first_numId = list_items[0]["numId"]
            first_ilvl = str(list_items[0]["ilvl"])
            fmt = numbering_map.get((first_numId, first_ilvl), "bullet")
            is_ordered = fmt in ("decimal", "lowerLetter", "lowerRoman",
                                 "upperLetter", "upperRoman")

            lvl = (2 if in_appendix else 1) + len(section_stack)

            def get_list_tag(numId, ilvl_int):
                fmt = numbering_map.get((numId, str(ilvl_int)), "bullet")
                ordered = fmt in ("decimal", "lowerLetter", "lowerRoman",
                                  "upperLetter", "upperRoman")
                return "orderedlist" if ordered else "itemizedlist"

            def render_list(items, base_lvl):
                """Render a list of items that may contain nested sub-lists."""
                if not items:
                    return
                tag = get_list_tag(items[0]["numId"], items[0]["ilvl"])
                xml_lines.append(f"{indent(base_lvl)}<{tag}>")

                idx = 0
                while idx < len(items):
                    li = items[idx]
                    content = render_inline(li["inline"])
                    xml_lines.append(f"{indent(base_lvl+1)}<listitem>")
                    xml_lines.append(f"{indent(base_lvl+2)}<para>{content}</para>")

                    # Collect any nested items that follow at a deeper level
                    nested = []
                    while (idx + 1 < len(items)
                           and items[idx + 1]["ilvl"] > li["ilvl"]):
                        idx += 1
                        nested.append(items[idx])

                    if nested:
                        render_list(nested, base_lvl + 2)

                    xml_lines.append(f"{indent(base_lvl+1)}</listitem>")
                    idx += 1

                xml_lines.append(f"{indent(base_lvl)}</{tag}>")

            render_list(list_items, lvl)
            continue

        # Handle Title (skip - already in articleinfo)
        if style == "Title":
            i += 1
            continue

        # Handle Normal paragraphs
        if style == "Normal":
            content = render_inline(elem["inline"])
            if not content.strip():
                i += 1
                continue

            lvl = (2 if in_appendix else 1) + len(section_stack)

            # Detect definition list paragraphs (bold term + tab + definition)
            # These appear in section 2 "Definitions and Acronyms" and in
            # schema description fields (Description:/Normative schema:/Identifier:)
            def is_definition_para(el):
                """Check if a paragraph looks like TERM<tab>Definition."""
                inl = el.get("inline", [])
                if not inl:
                    return False
                first = inl[0]
                if first["type"] != "text" or not first["bold"]:
                    return False
                # Check if raw text contains a tab (not stripped)
                raw = el.get("raw_text", "")
                return "\t" in raw

            if is_definition_para(elem):
                # Collect consecutive definition paragraphs
                def_items = []
                while i < len(elements) and elements[i]["style"] == "Normal" and is_definition_para(elements[i]):
                    def_items.append(elements[i])
                    i += 1

                xml_lines.append(f"{indent(lvl)}<variablelist>")
                for di in def_items:
                    raw = di.get("raw_text", "")
                    # Split on first tab
                    if "\t" in raw:
                        term, defn = raw.split("\t", 1)
                    else:
                        term = raw
                        defn = ""
                    term = term.strip()
                    defn = defn.strip()
                    # For entries with hyperlinks, render inline content minus the term
                    inl = di.get("inline", [])
                    defn_inline = []
                    found_tab = False
                    for item in inl:
                        if item["type"] == "text" and "\t" in item["text"] and not found_tab:
                            # Split on tab, keep what's after
                            after_tab = item["text"].split("\t", 1)[1]
                            if after_tab.strip():
                                defn_inline.append({"type": "text", "text": after_tab,
                                                    "bold": False, "italic": False})
                            found_tab = True
                            continue
                        if found_tab:
                            defn_inline.append(item)

                    if defn_inline:
                        defn_rendered = render_inline(defn_inline)
                    else:
                        defn_rendered = xml_escape(defn) if defn else ""

                    xml_lines.append(f"{indent(lvl+1)}<varlistentry>")
                    xml_lines.append(f'{indent(lvl+2)}<term>{xml_escape(term)}</term>')
                    xml_lines.append(f"{indent(lvl+2)}<listitem>")
                    if defn_rendered:
                        xml_lines.append(f"{indent(lvl+3)}<para>{defn_rendered}</para>")
                    else:
                        xml_lines.append(f"{indent(lvl+3)}<para></para>")
                    xml_lines.append(f"{indent(lvl+2)}</listitem>")
                    xml_lines.append(f"{indent(lvl+1)}</varlistentry>")
                xml_lines.append(f"{indent(lvl)}</variablelist>")
                continue

            # In bibliography sections, render references as bibliomixed
            if in_bibliography and bib_subsection and elem["inline"]:
                # Check if this is a reference entry (starts with [something] in bold)
                first = elem["inline"][0] if elem["inline"] else None
                if first and first["type"] == "text" and first["bold"] and first["text"].startswith("["):
                    # Extract the abbreviation
                    abbrev_match = re.match(r"\[([^\]]+)\]", first["text"])
                    if abbrev_match:
                        abbrev = abbrev_match.group(1)
                        bib_id = "BIB-" + slugify(abbrev)
                        xml_lines.append(f'{indent(lvl)}<bibliomixed id="{xml_escape(bib_id)}">')
                        xml_lines.append(f"{indent(lvl+1)}<abbrev>{xml_escape(abbrev)}</abbrev>")
                        # Render the rest
                        rest_content = render_inline(elem["inline"])
                        xml_lines.append(f"{indent(lvl+1)}{rest_content}")
                        xml_lines.append(f"{indent(lvl)}</bibliomixed>")
                        i += 1
                        continue

                # Non-reference paragraph in bibliography
                xml_lines.append(f"{indent(lvl)}<para>{content}</para>")
                i += 1
                continue

            xml_lines.append(f"{indent(lvl)}<para>{content}</para>")
            i += 1
            continue

        # Skip anything else
        i += 1

    # Close remaining sections
    if in_appendix:
        close_all_sections(xml_lines, base_indent=2)
        xml_lines.append("  </appendix>")
    else:
        close_all_sections(xml_lines)

    xml_lines.append("")
    xml_lines.append("</article>")

    # Write output
    output = "\n".join(xml_lines) + "\n"
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        f.write(output)

    print(f"Written {len(output)} bytes to {OUTPUT_PATH}")


if __name__ == "__main__":
    convert()

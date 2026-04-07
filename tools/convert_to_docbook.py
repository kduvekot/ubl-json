#!/usr/bin/env python3
"""Convert a UBL .docx specification to DocBook XML 4.5."""

import re
import docx
from lxml import etree
import zipfile

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W = f"{{{WNS}}}"
R = f"{{{RNS}}}"

DOCX_PATH = "source/UBL_2.5_JSON_Syntax_Binding_version_1.0_WD01.docx"
OUTPUT_PATH = "UBL-json.xml"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def slugify(text):
    """Create an uppercase slug from text for use in XML ids."""
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text.strip())
    return text.upper()


def xml_escape(text):
    """Escape special XML characters."""
    if text is None:
        return ""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


# ---------------------------------------------------------------------------
# Build hyperlink relationship map
# ---------------------------------------------------------------------------

def build_hyperlink_map(doc):
    """Map relationship IDs to URLs for external hyperlinks."""
    rels = doc.part.rels
    hmap = {}
    for rel_id, rel in rels.items():
        if "hyperlink" in str(rel.reltype).lower():
            target = rel.target_ref
            # Fix Word source erratum: common schema hyperlinks use "-2.json"
            # but should be "-2.5.json" (display text is correct, URL is not).
            # Safe when the .docx is corrected: "-2.5.json" won't match "-2.json".
            if target.startswith("json/schemas/common/") and target.endswith("-2.json"):
                target = target[:-len("-2.json")] + "-2.5.json"
            hmap[rel_id] = target
    return hmap


# ---------------------------------------------------------------------------
# Front-matter metadata parser
# ---------------------------------------------------------------------------

def _parse_person_entry(text):
    """Parse a person entry of the form:
      'Name (email), Organisation'
    or just 'Name (email),' (no org) or 'Name (email)'.

    Returns a dict with keys 'name', 'email', 'org' (org may be empty string).
    """
    text = text.strip().rstrip(",")
    # Match: Name (email), Org  OR  Name (email)
    m = re.match(r"^(.+?)\s*\(([^)]+)\)\s*(?:,\s*(.*))?$", text)
    if m:
        name = m.group(1).strip()
        email = m.group(2).strip()
        org = (m.group(3) or "").strip()
        return {"name": name, "email": email, "org": org}
    # Fallback: just store the whole string as name
    return {"name": text, "email": "", "org": ""}


def _normalise_label(text):
    """Strip leading/trailing whitespace and zero-width spaces from a label."""
    return text.replace("\u200b", "").strip().rstrip(":").strip()


def parse_front_matter(doc):
    """Parse front-matter metadata from the document paragraphs.

    Walks paragraphs from the beginning, treating:
      - The first Title-style paragraph as the document title.
      - The first two Subtitle-style paragraphs (before any label/value pairs)
        as status and date.
      - Subsequent Subtitle-style paragraphs as metadata labels, each followed
        by one or more Normal paragraphs as values.
    Stops at the first Heading 1 paragraph.

    Returns a dict with keys:
        title, status, date,
        this_version_urls, previous_version_urls, latest_version_urls,
        technical_committee,
        chairs   – list of {name, email, org}
        editors  – list of {name, email, org}
        related_work, abstract, citation_format, status_text
    Any key whose data was not found is omitted from the returned dict.
    """
    result = {}

    # Collect paragraphs up to (but not including) the first Heading 1.
    front_paras = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else "Normal"
        if style.startswith("Heading"):
            break
        front_paras.append(p)

    if not front_paras:
        return result

    idx = 0
    n = len(front_paras)

    # --- Title ---
    while idx < n:
        p = front_paras[idx]
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()
        if style == "Title" and text:
            result["title"] = text
            idx += 1
            break
        idx += 1

    # --- Status and Date: first two non-empty Subtitle paragraphs before any
    #     label/value section begins.  A label/value Subtitle is one that is
    #     *not* immediately preceded by another Subtitle (i.e. the first two
    #     consecutive Subtitles are the preamble status/date block).
    # Strategy: collect leading Subtitle paragraphs that form the preamble.
    # The preamble Subtitles are those that appear before the first Subtitle
    # whose *following* paragraph is Normal (i.e. a label with a value).
    # Simpler approach: the first Subtitle is status, the second is date,
    # and any further Subtitle paragraphs are labels.
    status_assigned = False
    date_assigned = False

    while idx < n:
        p = front_paras[idx]
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()

        if style == "Title":
            # Already handled above; skip any duplicates.
            idx += 1
            continue

        if style == "Subtitle":
            label = _normalise_label(text)
            if not status_assigned:
                result["status"] = label
                status_assigned = True
                idx += 1
                continue
            if not date_assigned:
                result["date"] = label
                date_assigned = True
                idx += 1
                continue
            # From here on, Subtitle paragraphs are labels with Normal values.
            break

        # Skip leading Normal (empty/decorative) paragraphs.
        idx += 1

    # --- Label / value sections ---
    # Mapping from normalised label text to result dict key and value handler.
    URL_LABELS = {
        "this version": "this_version_urls",
        "previous version": "previous_version_urls",
        "latest version": "latest_version_urls",
    }
    PERSON_LABELS = {"chair", "chairs", "editor", "editors"}
    LIST_LABELS = {"related work"}
    TEXT_LABELS = {
        "technical committee": "technical_committee",
        "abstract": "abstract",
        "citation format": "citation_format",
        "license, document status, and notices": "status_text",
        "license document status and notices": "status_text",
    }

    while idx < n:
        p = front_paras[idx]
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()

        if not text:
            idx += 1
            continue

        if style != "Subtitle":
            # A stray Normal paragraph not belonging to a label section; skip.
            idx += 1
            continue

        label = _normalise_label(text).lower()
        idx += 1

        # Collect the following Normal paragraphs as values until the next
        # Subtitle (or end of front matter).  Stop early if we encounter the
        # "Table of Contents" sentinel paragraph, which marks the end of the
        # meaningful front-matter content.
        values = []
        while idx < n:
            vp = front_paras[idx]
            vstyle = vp.style.name if vp.style else "Normal"
            vtext = vp.text.strip()
            if vstyle == "Subtitle":
                break
            if vtext == "Table of Contents":
                # Advance past this and any trailing empty paras, then stop
                # the outer label loop as well.
                while idx < n:
                    idx += 1
                n = idx  # force outer loop to terminate after this label
                break
            if vtext:
                values.append(vtext)
            idx += 1

        if not values:
            continue

        if label in URL_LABELS:
            key = URL_LABELS[label]
            # Each value line may contain a URL and an optional annotation like
            # "(Authoritative)".  Store as list of strings as-is.
            result[key] = values

        elif label in PERSON_LABELS:
            # Singular key: "chairs" / "editors"
            key = "chairs" if "chair" in label else "editors"
            people = []
            for v in values:
                # A value line ending with a bare comma (no org) is valid.
                people.append(_parse_person_entry(v))
            result[key] = people

        elif label in LIST_LABELS:
            result["related_work"] = values

        else:
            key = TEXT_LABELS.get(label)
            if key:
                result[key] = " ".join(values)
            else:
                # Unknown label — store under the normalised label as-is.
                safe_key = re.sub(r"[^a-z0-9]+", "_", label).strip("_")
                result[safe_key] = values if len(values) > 1 else values[0]

    return result


# ---------------------------------------------------------------------------
# Build numbering format map
# ---------------------------------------------------------------------------

def build_numbering_map(docx_path):
    """Build a map: (numId, ilvl) -> 'bullet' | 'decimal' | ... """
    with zipfile.ZipFile(docx_path) as z:
        with z.open("word/numbering.xml") as f:
            tree = etree.fromstring(f.read())

    ns = {"w": WNS}
    # abstractNumId -> { ilvl -> numFmt }
    abstract_map = {}
    for an in tree.findall(".//w:abstractNum", ns):
        aid = an.get(f"{W}abstractNumId")
        abstract_map[aid] = {}
        for lvl in an.findall(".//w:lvl", ns):
            ilvl = lvl.get(f"{W}ilvl")
            nf = lvl.find(".//w:numFmt", ns)
            fmt = nf.get(f"{W}val") if nf is not None else "bullet"
            abstract_map[aid][ilvl] = fmt

    # numId -> abstractNumId
    num_to_abstract = {}
    for num in tree.findall(".//w:num", ns):
        nid = num.get(f"{W}numId")
        anr = num.find(".//w:abstractNumId", ns)
        if anr is not None:
            num_to_abstract[nid] = anr.get(f"{W}val")

    # (numId, ilvl) -> fmt
    result = {}
    for nid, aid in num_to_abstract.items():
        if aid in abstract_map:
            for ilvl, fmt in abstract_map[aid].items():
                result[(nid, ilvl)] = fmt
    return result


# ---------------------------------------------------------------------------
# Inline content extraction
# ---------------------------------------------------------------------------

def extract_inline_content(paragraph, hyperlink_map):
    """Walk paragraph XML to extract runs and hyperlinks in document order.

    Returns a list of dicts:
      {'type': 'text', 'text': str, 'bold': bool, 'italic': bool}
      {'type': 'hyperlink', 'url': str, 'text': str}
    """
    result = []
    p_elem = paragraph._element

    for child in p_elem:
        tag = child.tag
        if tag == f"{W}r":
            # Collect text from <w:t> and treat <w:tab/> as a tab character
            text_parts = []
            for sub in child:
                if sub.tag == f"{W}t" and sub.text:
                    text_parts.append(sub.text)
                elif sub.tag == f"{W}tab":
                    text_parts.append("\t")
            text = "".join(text_parts)
            if not text:
                continue
            rPr = child.find(f"{W}rPr")
            bold = False
            italic = False
            monospace = False
            if rPr is not None:
                b_elem = rPr.find(f"{W}b")
                if b_elem is not None:
                    val = b_elem.get(f"{W}val")
                    bold = val != "0" if val is not None else True
                i_elem = rPr.find(f"{W}i")
                if i_elem is not None:
                    val = i_elem.get(f"{W}val")
                    italic = val != "0" if val is not None else True
                # Check for monospace character style
                rStyle = rPr.find(f"{W}rStyle")
                if rStyle is not None:
                    style_val = rStyle.get(f"{W}val", "")
                    if "monospace" in style_val.lower() or "code" in style_val.lower():
                        monospace = True
            result.append({"type": "text", "text": text, "bold": bold, "italic": italic, "monospace": monospace})

        elif tag == f"{W}hyperlink":
            rid = child.get(f"{R}id")
            texts = []
            for r in child.findall(f".//{W}t"):
                if r.text:
                    texts.append(r.text)
            link_text = "".join(texts)
            url = hyperlink_map.get(rid, "") if rid else ""
            if link_text:
                result.append({"type": "hyperlink", "url": url, "text": link_text})

    # Merge adjacent runs with identical formatting (Word often splits a single
    # word like "ExtensionURI" into multiple runs, e.g. "E" + "xtensionURI").
    merged = []
    for item in result:
        if (merged
                and item["type"] == "text"
                and merged[-1]["type"] == "text"
                and item["bold"] == merged[-1]["bold"]
                and item["italic"] == merged[-1]["italic"]
                and item.get("monospace") == merged[-1].get("monospace")):
            merged[-1]["text"] += item["text"]
        else:
            merged.append(item)
    return merged


def render_inline(items):
    """Render a list of inline items to DocBook XML string."""
    parts = []
    for item in items:
        if item["type"] == "hyperlink":
            url = xml_escape(item["url"])
            text = xml_escape(item["text"])
            if url:
                parts.append(f'<ulink url="{url}">{text}</ulink>')
            else:
                parts.append(text)
        else:
            # Replace tabs with spaces for inline text (tabs are structural
            # separators in the .docx, not meaningful in DocBook inline content)
            text = xml_escape(item["text"].replace("\t", " "))
            if item.get("monospace"):
                parts.append(f'<literal moreinfo="none">{text}</literal>')
            elif item["bold"] and item["italic"]:
                parts.append(f'<emphasis role="bold"><emphasis>{text}</emphasis></emphasis>')
            elif item["bold"]:
                parts.append(f'<emphasis role="bold">{text}</emphasis>')
            elif item["italic"]:
                parts.append(f'<emphasis>{text}</emphasis>')
            else:
                parts.append(text)
    return "".join(parts)


# ---------------------------------------------------------------------------
# Paragraph info extraction
# ---------------------------------------------------------------------------

def get_paragraph_info(p):
    """Extract style, list info from a paragraph."""
    style = p.style.name if p.style else "Normal"
    numId = None
    ilvl = None
    pPr = p._element.find(f"{W}pPr")
    if pPr is not None:
        numPr = pPr.find(f"{W}numPr")
        if numPr is not None:
            il = numPr.find(f"{W}ilvl")
            ni = numPr.find(f"{W}numId")
            if il is not None:
                ilvl = il.get(f"{W}val")
            if ni is not None:
                numId = ni.get(f"{W}val")
    return style, numId, ilvl


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def render_table(table, hyperlink_map, title=None):
    """Render a docx table to DocBook CALS table XML.

    title: optional string to use as the table title.  When omitted the title
    is derived from the text of the first cell of the first header row.
    """
    rows = table.rows
    if not rows:
        return ""

    if title is None:
        # Derive from the first cell's text (strip trailing plural marker etc.)
        title = rows[0].cells[0].text.strip() if rows[0].cells else "Table"

    table_id = f"T-{slugify(title)}"
    ncols = len(table.columns)
    lines = []
    lines.append(f'<table id="{table_id}" frame="all">')
    lines.append(f'  <title>{xml_escape(title)}</title>')
    lines.append(f'  <tgroup cols="{ncols}">')
    for ci in range(ncols):
        lines.append(f'    <colspec colnum="{ci+1}" colname="col{ci+1}"/>')

    def render_cell(cell):
        """Render a table cell with inline formatting preserved."""
        parts = []
        for para in cell.paragraphs:
            inline = extract_inline_content(para, hyperlink_map)
            if inline:
                parts.append(render_inline(inline))
        return " ".join(parts) if parts else xml_escape(cell.text.strip())

    # First row is header
    lines.append("    <thead>")
    lines.append("      <row>")
    for cell in rows[0].cells:
        lines.append(f"        <entry>{render_cell(cell)}</entry>")
    lines.append("      </row>")
    lines.append("    </thead>")

    lines.append("    <tbody>")
    for row in rows[1:]:
        lines.append("      <row>")
        for cell in row.cells:
            lines.append(f"        <entry>{render_cell(cell)}</entry>")
        lines.append("      </row>")
    lines.append("    </tbody>")
    lines.append("  </tgroup>")
    lines.append("</table>")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Table position detection
# ---------------------------------------------------------------------------

def find_table_positions(doc):
    """Return a dict mapping paragraph_index -> table_object_index.

    paragraph_index is the 0-based index (in doc.paragraphs) of the paragraph
    *after* which the corresponding table appears in the document body.
    table_object_index is the 0-based index into doc.tables.

    Works by walking the direct children of the document body and counting
    <w:p> elements; each time a <w:tbl> is encountered the most recently seen
    paragraph index is recorded.
    """
    table_tag = f"{W}tbl"
    para_tag = f"{W}p"

    positions = {}
    para_count = 0   # running count of <w:p> children seen so far
    table_idx = 0    # index into doc.tables

    for child in doc.element.body:
        if child.tag == para_tag:
            para_count += 1
        elif child.tag == table_tag:
            # The table belongs after paragraph index para_count - 1
            positions[para_count - 1] = table_idx
            table_idx += 1

    return positions


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert(docx_path=DOCX_PATH, output_path=OUTPUT_PATH):
    doc = docx.Document(docx_path)
    hyperlink_map = build_hyperlink_map(doc)
    numbering_map = build_numbering_map(docx_path)

    # -----------------------------------------------------------------------
    # Phase 1: Parse all paragraphs into structured elements
    # -----------------------------------------------------------------------
    elements = []  # list of dicts describing each document element

    # Build a mapping of paragraph_index -> table_object_index so we know
    # where to insert table markers as we walk doc.paragraphs.
    table_positions = find_table_positions(doc)

    seen_first_heading = False

    for i, p in enumerate(doc.paragraphs):
        style, numId, ilvl = get_paragraph_info(p)

        # Skip ToC entries
        if style.startswith("toc"):
            continue

        # Skip the "Table of Contents" heading
        if style == "Heading 1" and p.text.strip() == "Table of Contents":
            continue

        # Track whether we've seen the first Heading 1.
        # Paragraphs before it are front matter (already handled by
        # parse_front_matter()) and must not be emitted as body content.
        if style == "Heading 1":
            seen_first_heading = True
        elif not seen_first_heading:
            # Front-matter region: skip Title, Subtitle, Normal, etc.
            # They are already processed by parse_front_matter().
            if i in table_positions:
                pass  # ignore any table that appears before the first heading
            continue

        # Handle Subtitle paragraphs in the body (unusual, but treat as Normal)
        if style == "Subtitle":
            style = "Normal"

        text = p.text.strip()
        raw_text = p.text  # preserve tabs for definition detection

        inline = extract_inline_content(p, hyperlink_map)

        # Skip truly empty Normal paragraphs
        if not text and not inline and style == "Normal":
            # Still insert a TABLE marker if a table follows this empty paragraph.
            if i in table_positions:
                elements.append({"style": "TABLE", "index": table_positions[i]})
            continue

        elem = {
            "index": i,
            "style": style,
            "text": text,
            "raw_text": raw_text,
            "inline": inline,
            "numId": numId,
            "ilvl": ilvl,
        }
        elements.append(elem)

        # Insert a TABLE marker if a table follows this paragraph in the body.
        if i in table_positions:
            elements.append({"style": "TABLE", "index": table_positions[i]})

    # -----------------------------------------------------------------------
    # Phase 2: Group elements into a tree structure
    # -----------------------------------------------------------------------
    # We'll generate the XML by walking through elements and tracking
    # the section nesting stack.

    xml_lines = []

    def indent(level):
        return "  " * level

    # --- Parse front-matter metadata ---
    meta = parse_front_matter(doc)

    # -----------------------------------------------------------------------
    # Build entity definitions from parsed metadata
    # -----------------------------------------------------------------------

    def _url_base(url):
        """Strip filename (last path segment) and trailing slash from a URL."""
        url = url.strip()
        # Remove any trailing annotation like " (Authoritative)"
        url = re.sub(r"\s*\([^)]*\)\s*$", "", url).strip()
        # Strip the last path component (the filename)
        return url.rsplit("/", 1)[0].rstrip("/")

    # this-loc: base URL from first "this version" URL
    this_version_urls = meta.get("this_version_urls", [])
    this_loc = _url_base(this_version_urls[0]) if this_version_urls else ""

    # latest-loc: base URL from first "latest version" URL
    latest_version_urls = meta.get("latest_version_urls", [])
    latest_loc = _url_base(latest_version_urls[0]) if latest_version_urls else ""

    # previous-loc: base URL from first "previous version" URL (may be empty)
    previous_version_urls = meta.get("previous_version_urls", [])
    previous_loc = _url_base(previous_version_urls[0]) if previous_version_urls else ""

    # stage: extract from the last path segment of this-loc
    # e.g. "https://docs.oasis-open.org/ubl/csd01-UBL-2.5-JSON-1.0" -> "csd01"
    stage = ""
    if this_loc:
        last_seg = this_loc.rsplit("/", 1)[-1]
        m = re.match(r"^([a-z]+\d+)-", last_seg)
        if m:
            stage = m.group(1)

    # spec-version: parse from title "Version X.Y"
    spec_version = ""
    title_str = meta.get("title", "")
    m = re.search(r"[Vv]ersion\s+([\d.]+)", title_str)
    if m:
        spec_version = m.group(1)

    entities = {
        "name": "UBL",
        "version": "2.5",
        "spec-version": spec_version,
        "stage": stage,
        "standard": meta.get("status", ""),
        "pubdate": meta.get("date", ""),
        "title": title_str,
        "this-loc": this_loc,
        "latest-loc": latest_loc,
        "previous-loc": previous_loc,
        "committee": meta.get("technical_committee", ""),
        "abstract-text": meta.get("abstract", ""),
    }

    # -----------------------------------------------------------------------
    # XML declaration and DOCTYPE with entity declarations
    # -----------------------------------------------------------------------
    xml_lines.append('<?xml version="1.0" encoding="UTF-8"?>')
    xml_lines.append('<!DOCTYPE article PUBLIC "-//OASIS//DTD DocBook XML V4.5//EN"')
    xml_lines.append('  "http://www.oasis-open.org/docbook/xml/4.5/docbookx.dtd" [')
    for ent_name, ent_value in entities.items():
        xml_lines.append(f'  <!ENTITY {ent_name} "{xml_escape(ent_value)}">')
    xml_lines.append(']>')
    xml_lines.append("")

    # -----------------------------------------------------------------------
    # Generate <articleinfo> using entity references
    # -----------------------------------------------------------------------
    xml_lines.append('<article status="&standard;" lang="en">')
    xml_lines.append("  <articleinfo>")

    # Product name / number (always UBL 2.5 for this document)
    xml_lines.append('    <productname class="trade">&name;</productname>')
    xml_lines.append('    <productnumber>&version;</productnumber>')

    # Standards track (invariant — not version-specific)
    xml_lines.append('    <releaseinfo role="track">Standards Track Work Product</releaseinfo>')

    # "This version" URLs — HTML first, then PDF, then XML (authoritative)
    xml_lines.append('    <releaseinfo role="OASIS-specification-this">&this-loc;/UBL-2.5-JSON-&spec-version;.html</releaseinfo>')
    xml_lines.append('    <releaseinfo role="OASIS-specification-this">&this-loc;/UBL-2.5-JSON-&spec-version;.pdf</releaseinfo>')
    xml_lines.append('    <releaseinfo role="OASIS-specification-this-authoritative">&this-loc;/UBL-2.5-JSON-&spec-version;.xml</releaseinfo>')

    # "Previous version" URLs (if available)
    if meta.get("previous_version_urls"):
        xml_lines.append('    <releaseinfo role="OASIS-specification-previous">&previous-loc;/UBL-2.5-JSON-&spec-version;.html</releaseinfo>')
        xml_lines.append('    <releaseinfo role="OASIS-specification-previous">&previous-loc;/UBL-2.5-JSON-&spec-version;.pdf</releaseinfo>')
        xml_lines.append('    <releaseinfo role="OASIS-specification-previous-authoritative">&previous-loc;/UBL-2.5-JSON-&spec-version;.xml</releaseinfo>')

    # "Latest version" URLs
    xml_lines.append('    <releaseinfo role="OASIS-specification-latest">&latest-loc;/UBL-2.5-JSON-&spec-version;.html</releaseinfo>')
    xml_lines.append('    <releaseinfo role="OASIS-specification-latest">&latest-loc;/UBL-2.5-JSON-&spec-version;.pdf</releaseinfo>')

    # Title
    xml_lines.append("    <title>&title;</title>")

    # Technical committee
    if meta.get("technical_committee"):
        xml_lines.append('    <releaseinfo role="committee">&committee;</releaseinfo>')

    # Editors (structured content — kept dynamic, not entities)
    editors = meta.get("editors", [])
    if editors:
        xml_lines.append("    <authorgroup>")
        for ed in editors:
            xml_lines.append("      <editor>")
            name = ed.get("name", "")
            # Split name into firstname / surname on the last space
            if " " in name:
                firstname, surname = name.rsplit(" ", 1)
            else:
                firstname, surname = "", name
            xml_lines.append(f"        <firstname>{xml_escape(firstname)}</firstname>")
            xml_lines.append(f"        <surname>{xml_escape(surname)}</surname>")
            org = ed.get("org", "")
            if org:
                xml_lines.append(f"        <affiliation><orgname>{xml_escape(org)}</orgname></affiliation>")
            email = ed.get("email", "")
            if email:
                xml_lines.append(f"        <email>{xml_escape(email)}</email>")
            xml_lines.append("      </editor>")
        xml_lines.append("    </authorgroup>")

    # Publication date
    if meta.get("date"):
        xml_lines.append("    <pubdate>&pubdate;</pubdate>")

    # Related work (structured content — kept dynamic, not entities)
    related = meta.get("related_work", [])
    if related:
        xml_lines.append('    <legalnotice role="related">')
        xml_lines.append("      <title>Related work</title>")
        # If the first line is an introductory sentence, emit it as a <para>
        # then the remainder as <bibliomixed> entries inside a <bibliolist>.
        intro_lines = []
        bib_entries = []
        for line in related:
            # Heuristic: lines starting with '[' are bibliography references
            if line.lstrip().startswith("["):
                bib_entries.append(line)
            else:
                intro_lines.append(line)
        for intro in intro_lines:
            xml_lines.append(f"      <para>{xml_escape(intro)}</para>")
        if bib_entries:
            xml_lines.append("      <bibliolist>")
            for entry in bib_entries:
                xml_lines.append(f"        <bibliomixed>{xml_escape(entry)}</bibliomixed>")
            xml_lines.append("      </bibliolist>")
        xml_lines.append("    </legalnotice>")

    # Abstract
    if meta.get("abstract"):
        xml_lines.append("    <abstract>")
        xml_lines.append("      <para>&abstract-text;</para>")
        xml_lines.append("    </abstract>")

    # Citation format — build structured citation from metadata
    xml_lines.append('    <legalnotice role="citation" id="CITATION">')
    xml_lines.append("      <title>Citation format</title>")
    xml_lines.append("      <para>When referencing this specification the following citation format should be used:</para>")
    xml_lines.append('      <bibliolist id="citationfmt">')
    xml_lines.append('        <bibliomixed id="UBL-JSON" conformance="skip">')
    xml_lines.append('          <abbrev condition="oasis">UBL-2.5-JSON-&spec-version;</abbrev>')
    xml_lines.append('          <citetitle>&title;.</citetitle>')

    # Build "Edited by X, Y and Z." from editors
    editors = meta.get("editors", [])
    if editors:
        editor_names = []
        for ed in editors:
            name = ed.get("name", "")
            editor_names.append(name)
        if len(editor_names) == 1:
            editors_str = editor_names[0]
        elif len(editor_names) == 2:
            editors_str = f"{editor_names[0]} and {editor_names[1]}"
        else:
            editors_str = ", ".join(editor_names[:-1]) + f" and {editor_names[-1]}"
        xml_lines.append(f'          <bibliomisc>Edited by {xml_escape(editors_str)}.</bibliomisc>')

    xml_lines.append('          <date>&pubdate;.</date>')
    xml_lines.append('          <releaseinfo>&standard;.</releaseinfo>')
    xml_lines.append('          <bibliomisc>')
    xml_lines.append('            <ulink url="&this-loc;/UBL-2.5-JSON-&spec-version;.html">&this-loc;/UBL-2.5-JSON-&spec-version;.html</ulink>.')
    xml_lines.append('            Latest stage: <ulink url="&latest-loc;/UBL-2.5-JSON-&spec-version;.html">&latest-loc;/UBL-2.5-JSON-&spec-version;.html</ulink>.')
    xml_lines.append('          </bibliomisc>')
    xml_lines.append('        </bibliomixed>')
    xml_lines.append('      </bibliolist>')
    xml_lines.append('    </legalnotice>')

    # Status / notices text (copyright / license notice)
    status_text = meta.get("status_text", "")
    if status_text:
        xml_lines.append('    <legalnotice role="notices">')
        xml_lines.append("      <title>Notices</title>")
        xml_lines.append(f"      <para>{xml_escape(status_text)}</para>")
        xml_lines.append("    </legalnotice>")

    xml_lines.append("  </articleinfo>")
    xml_lines.append("")

    # -----------------------------------------------------------------------
    # State machine for generating nested sections
    # -----------------------------------------------------------------------
    heading_level_map = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
    }

    # Section stack: list of heading levels currently open
    section_stack = []
    # Track parent section titles for disambiguating IDs
    section_title_stack = []
    in_bibliography = False
    bib_subsection = None  # "normative" or "informative"
    in_appendix = False
    appendix_name = None
    used_ids = set()

    def close_sections_to(target_level, lines, base_indent=1):
        """Close sections until we're at a level where we can open target_level."""
        while section_stack and section_stack[-1] >= target_level:
            section_stack.pop()
            lvl = base_indent + len(section_stack)
            lines.append(f"{indent(lvl)}</section>")

    def close_all_sections(lines, base_indent=1):
        while section_stack:
            section_stack.pop()
            lvl = base_indent + len(section_stack)
            lines.append(f"{indent(lvl)}</section>")

    # Process elements
    i = 0
    while i < len(elements):
        elem = elements[i]
        style = elem["style"]

        # Handle table insertion marker
        if style == "TABLE":
            tbl = doc.tables[elem["index"]]
            base = 1 + len(section_stack)
            table_xml = render_table(tbl, hyperlink_map)
            for line in table_xml.split("\n"):
                xml_lines.append(f"{indent(base)}{line}")
            i += 1
            continue

        text = elem.get("text", "")

        # Detect any Annex (A, B, C, ...) or Appendix (1, 2, ...) generically
        annex_match = re.match(r"Annex\s+([A-Z])", text)
        appendix_match = re.match(r"Appendix\s+(\w+)", text)
        if style == "Heading 1" and (annex_match or appendix_match):
            annex_letter = annex_match.group(1) if annex_match else appendix_match.group(1)
            # Close previous appendix or open sections
            if in_appendix:
                close_all_sections(xml_lines, base_indent=2)
                xml_lines.append("  </appendix>")
                in_appendix = False
            else:
                close_all_sections(xml_lines)

            # Build an id slug from the part of the title after "Annex/Appendix X"
            # The descriptive title (without the "Annex A" / "Appendix 1" prefix)
            # is used as the <title> — DocBook auto-generates appendix labels.
            parts = text.split(None, 2)
            descriptive_title = parts[2] if len(parts) > 2 else text
            slug = slugify(descriptive_title)

            xml_lines.append("")
            xml_lines.append(f'  <appendix id="A-{slug}">')
            xml_lines.append(f"    <title>{xml_escape(descriptive_title)}</title>")
            in_appendix = True
            appendix_name = annex_letter
            section_stack.clear()
            section_title_stack.clear()

            # Annex B is the bibliography section
            if annex_match and annex_letter == "B":
                in_bibliography = True
            else:
                in_bibliography = False

            i += 1
            continue

        # Handle headings
        if style in heading_level_map:
            hlevel = heading_level_map[style]

            if in_appendix:
                base = 2
            else:
                base = 1

            # Close sections as needed
            close_sections_to(hlevel, xml_lines, base_indent=base)

            # Also trim the section_title_stack to match
            while len(section_title_stack) >= hlevel:
                section_title_stack.pop()

            # Generate a unique section ID
            slug = slugify(text)
            if not slug:
                slug = "PLACEHOLDER"
            section_id = "S-" + slug
            # If the ID is already used, prefix with parent section title
            if section_id in used_ids:
                parent = section_title_stack[-1] if section_title_stack else ""
                section_id = "S-" + slugify(parent) + "-" + slug
            # Final dedup with numeric suffix if still colliding
            base_id = section_id
            counter = 2
            while section_id in used_ids:
                section_id = f"{base_id}-{counter}"
                counter += 1
            used_ids.add(section_id)

            section_title_stack.append(text)

            # Special handling for Annex B subsections
            if in_bibliography and hlevel == 2:
                if "Normative" in text:
                    bib_subsection = "normative"
                elif "Informative" in text:
                    bib_subsection = "informative"

            lvl = base + len(section_stack)
            xml_lines.append(f"")
            xml_lines.append(f'{indent(lvl)}<section id="{xml_escape(section_id)}">')
            xml_lines.append(f"{indent(lvl+1)}<title>{xml_escape(text)}</title>")
            section_stack.append(hlevel)
            i += 1
            continue

        # Handle Monospace (code blocks) - collect consecutive ones
        if style == "Monospace":
            code_lines = []
            lang = None
            while i < len(elements) and elements[i]["style"] == "Monospace":
                line = elements[i]["text"]
                stripped = line.strip()
                # Detect opening fence like ```json or ```xml
                fence_match = re.match(r"^```(\w+)?$", stripped)
                if fence_match:
                    if fence_match.group(1):
                        lang = fence_match.group(1)
                    # Skip both opening and closing fences
                    i += 1
                    continue
                code_lines.append(line)
                i += 1
            code_text = "\n".join(code_lines)
            lvl = (2 if in_appendix else 1) + len(section_stack)
            lang_attr = f' language="{xml_escape(lang)}"' if lang else ""
            xml_lines.append(f"{indent(lvl)}<programlisting{lang_attr}><![CDATA[{code_text}]]></programlisting>")
            continue

        # Handle List Paragraphs - collect consecutive ones into a list
        if style == "List Paragraph":
            list_items = []
            while i < len(elements) and elements[i]["style"] == "List Paragraph":
                e = elements[i]
                list_items.append({
                    "text": e["text"],
                    "inline": e["inline"],
                    "numId": e["numId"],
                    "ilvl": int(e["ilvl"] or "0"),
                })
                i += 1

            # Determine if this is an ordered or itemized list
            first_numId = list_items[0]["numId"]
            first_ilvl = str(list_items[0]["ilvl"])
            fmt = numbering_map.get((first_numId, first_ilvl), "bullet")
            is_ordered = fmt in ("decimal", "lowerLetter", "lowerRoman",
                                 "upperLetter", "upperRoman")

            lvl = (2 if in_appendix else 1) + len(section_stack)

            def get_list_tag(numId, ilvl_int):
                fmt = numbering_map.get((numId, str(ilvl_int)), "bullet")
                ordered = fmt in ("decimal", "lowerLetter", "lowerRoman",
                                  "upperLetter", "upperRoman")
                return "orderedlist" if ordered else "itemizedlist"

            def render_list(items, base_lvl):
                """Render a list of items that may contain nested sub-lists."""
                if not items:
                    return
                tag = get_list_tag(items[0]["numId"], items[0]["ilvl"])
                xml_lines.append(f"{indent(base_lvl)}<{tag}>")

                idx = 0
                while idx < len(items):
                    li = items[idx]
                    content = render_inline(li["inline"])
                    xml_lines.append(f"{indent(base_lvl+1)}<listitem>")
                    xml_lines.append(f"{indent(base_lvl+2)}<para>{content}</para>")

                    # Collect any nested items that follow at a deeper level
                    nested = []
                    while (idx + 1 < len(items)
                           and items[idx + 1]["ilvl"] > li["ilvl"]):
                        idx += 1
                        nested.append(items[idx])

                    if nested:
                        render_list(nested, base_lvl + 2)

                    xml_lines.append(f"{indent(base_lvl+1)}</listitem>")
                    idx += 1

                xml_lines.append(f"{indent(base_lvl)}</{tag}>")

            render_list(list_items, lvl)
            continue

        # Handle Title (skip - already in articleinfo)
        if style == "Title":
            i += 1
            continue

        # Handle Normal paragraphs
        if style == "Normal":
            content = render_inline(elem["inline"])
            if not content.strip():
                i += 1
                continue

            lvl = (2 if in_appendix else 1) + len(section_stack)

            # Detect definition list paragraphs (bold term + tab + definition)
            # These appear in section 2 "Definitions and Acronyms" and in
            # schema description fields (Description:/Normative schema:/Identifier:)
            def is_definition_para(el):
                """Check if a paragraph looks like TERM<tab>Definition."""
                inl = el.get("inline", [])
                if not inl:
                    return False
                first = inl[0]
                if first["type"] != "text" or not first["bold"]:
                    return False
                # Check if raw text contains a tab (not stripped)
                raw = el.get("raw_text", "")
                return "\t" in raw

            if is_definition_para(elem):
                # Collect consecutive definition paragraphs
                def_items = []
                while i < len(elements) and elements[i]["style"] == "Normal" and is_definition_para(elements[i]):
                    def_items.append(elements[i])
                    i += 1

                xml_lines.append(f"{indent(lvl)}<variablelist>")
                for di in def_items:
                    raw = di.get("raw_text", "")
                    # Split on first tab
                    if "\t" in raw:
                        term, defn = raw.split("\t", 1)
                    else:
                        term = raw
                        defn = ""
                    term = term.strip()
                    defn = defn.strip()
                    # Split inline content at the tab: items before it form the
                    # term, items after it form the definition.
                    inl = di.get("inline", [])
                    term_inline = []
                    defn_inline = []
                    found_tab = False
                    for item in inl:
                        if item["type"] == "text" and "\t" in item["text"] and not found_tab:
                            before_tab, after_tab = item["text"].split("\t", 1)
                            if before_tab.strip():
                                term_inline.append({**item, "text": before_tab})
                            if after_tab.strip():
                                defn_inline.append({**item, "text": after_tab})
                            found_tab = True
                            continue
                        if found_tab:
                            defn_inline.append(item)
                        else:
                            term_inline.append(item)

                    if defn_inline:
                        defn_rendered = render_inline(defn_inline)
                    else:
                        defn_rendered = xml_escape(defn) if defn else ""

                    term_rendered = render_inline(term_inline) if term_inline else xml_escape(term)

                    xml_lines.append(f"{indent(lvl+1)}<varlistentry>")
                    xml_lines.append(f'{indent(lvl+2)}<term>{term_rendered}</term>')
                    xml_lines.append(f"{indent(lvl+2)}<listitem>")
                    if defn_rendered:
                        xml_lines.append(f"{indent(lvl+3)}<para>{defn_rendered}</para>")
                    else:
                        xml_lines.append(f"{indent(lvl+3)}<para></para>")
                    xml_lines.append(f"{indent(lvl+2)}</listitem>")
                    xml_lines.append(f"{indent(lvl+1)}</varlistentry>")
                xml_lines.append(f"{indent(lvl)}</variablelist>")
                continue

            # In bibliography sections, render references as bibliomixed
            if in_bibliography and bib_subsection and elem["inline"]:
                # Check if this is a reference entry (starts with [something] in bold)
                first = elem["inline"][0] if elem["inline"] else None
                if first and first["type"] == "text" and first["bold"] and first["text"].startswith("["):
                    # Extract the abbreviation
                    abbrev_match = re.match(r"\[([^\]]+)\]", first["text"])
                    if abbrev_match:
                        abbrev = abbrev_match.group(1)
                        bib_id = "BIB-" + slugify(abbrev)
                        xml_lines.append(f'{indent(lvl)}<bibliomixed id="{xml_escape(bib_id)}">')
                        xml_lines.append(f"{indent(lvl+1)}<abbrev>{xml_escape(abbrev)}</abbrev>")
                        # Render the rest
                        rest_content = render_inline(elem["inline"])
                        xml_lines.append(f"{indent(lvl+1)}{rest_content}")
                        xml_lines.append(f"{indent(lvl)}</bibliomixed>")
                        i += 1
                        continue

                # Non-reference paragraph in bibliography
                xml_lines.append(f"{indent(lvl)}<para>{content}</para>")
                i += 1
                continue

            xml_lines.append(f"{indent(lvl)}<para>{content}</para>")
            i += 1
            continue

        # Skip anything else
        i += 1

    # Close remaining sections
    if in_appendix:
        close_all_sections(xml_lines, base_indent=2)
        xml_lines.append("  </appendix>")
    else:
        close_all_sections(xml_lines)

    xml_lines.append("")
    xml_lines.append("</article>")

    # Write output
    output = "\n".join(xml_lines) + "\n"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(output)

    print(f"Written {len(output)} bytes to {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(
        description="Convert a UBL .docx specification to DocBook XML 4.5"
    )
    parser.add_argument(
        "input", nargs="?", default=DOCX_PATH,
        help=f"Path to the .docx file (default: {DOCX_PATH})"
    )
    parser.add_argument(
        "-o", "--output",
        help="Output XML path (default: derived from input filename)"
    )
    args = parser.parse_args()

    docx_path = args.input
    if args.output:
        output_path = args.output
    elif docx_path == DOCX_PATH:
        # Use the hardcoded output path when running with the default input
        output_path = OUTPUT_PATH
    else:
        # Derive output from input: strip path, change extension
        import os
        base = os.path.splitext(os.path.basename(docx_path))[0]
        # Convert underscores/spaces to hyphens for the output filename
        base = base.replace("_", "-").replace(" ", "-")
        output_path = base + ".xml"

    convert(docx_path, output_path)
